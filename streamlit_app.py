import os
import streamlit as st
import pdfplumber
import docx
import sqlite3
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document

# -------------------------------
# Database Setup (SQLite)
# -------------------------------
conn = sqlite3.connect("resumes.db")
c = conn.cursor()
c.execute("""CREATE TABLE IF NOT EXISTS resumes (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT,
    job_role TEXT,
    ats_score REAL,
    match_score REAL,
    recommendations TEXT,
    interview_questions TEXT,
    salary_prediction TEXT
)""")
conn.commit()

# -------------------------------
# Font handling
# -------------------------------
# Build an absolute path to the font file based on this script's location,
# so it works no matter what directory Streamlit is launched from.
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FONT_PATH = os.path.join(BASE_DIR, "DejaVuSans.ttf")
UNICODE_FONT_AVAILABLE = os.path.isfile(FONT_PATH)

# -------------------------------
# Helper Functions
# -------------------------------
def extract_text_from_pdf(uploaded_file):
    text = ""
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

def calculate_ats_score(resume_text, job_description):
    resume_words = set(resume_text.lower().split())
    jd_words = set(job_description.lower().split())
    if len(jd_words) == 0:
        return 0
    match = resume_words.intersection(jd_words)
    return round(len(match) / len(jd_words) * 100, 2)

def generate_recommendations(job_role):
    role_lower = job_role.lower()
    role_data = [
        (["data analyst", "data analytics", "business intelligence", "bi analyst"],
         ["Learn SQL", "Practice Tableau/Power BI dashboards", "Add Python data projects"]),
        (["software engineer", "software developer", "backend", "full stack", "sde"],
         ["Strengthen DSA", "Contribute to GitHub projects", "Add cloud experience"]),
        (["ai engineer", "machine learning", "ml engineer", "data scientist"],
         ["Build ML models", "Add NLP/CV projects", "Learn TensorFlow/PyTorch"]),
        (["web developer", "frontend", "front-end", "react", "javascript developer"],
         ["Build responsive UI projects", "Learn React/Vue in depth", "Add a strong portfolio site"]),
        (["mobile app developer", "android developer", "ios developer", "flutter"],
         ["Publish an app on Play Store/App Store", "Learn Kotlin/Swift/Flutter deeply", "Add UI/UX polish to projects"]),
        (["marketing", "seo", "content", "digital marketing"],
         ["Learn SEO and analytics tools", "Build a content portfolio", "Get certified in Google Analytics"]),
        (["hr", "human resource", "recruit", "talent"],
         ["Learn HR software (ATS/HRMS)", "Build strong communication examples", "Get familiar with labor law basics"]),
        (["devops", "cloud engineer", "site reliability", "sre"],
         ["Learn Docker/Kubernetes", "Get an AWS/Azure certification", "Build CI/CD pipeline projects"]),
        (["cybersecurity", "security analyst", "ethical hacker", "penetration tester"],
         ["Get a Security+ or CEH certification", "Practice on TryHackMe/HackTheBox", "Build a home lab for network security"]),
        (["qa", "quality assurance", "software tester", "test engineer"],
         ["Learn automation tools like Selenium", "Understand test case design", "Add bug-tracking tool experience (Jira)"]),
        (["business analyst", "product analyst"],
         ["Learn SQL and Excel deeply", "Practice writing user stories", "Add process improvement projects"]),
        (["project manager", "program manager", "scrum master"],
         ["Get a PMP/Scrum certification", "Add examples of managing timelines/budgets", "Learn tools like Jira/Trello/MS Project"]),
        (["accountant", "finance", "financial analyst"],
         ["Get certified in Tally/Excel/SAP", "Add examples of financial reporting", "Understand GST/taxation basics"]),
        (["sales", "business development"],
         ["Add measurable sales targets achieved", "Highlight CRM tool experience", "Practice negotiation and pitching skills"]),
        (["graphic designer", "ui/ux designer", "ux designer", "product designer"],
         ["Build a strong visual portfolio (Behance/Dribbble)", "Learn Figma/Adobe XD deeply", "Add case studies explaining design decisions"]),
        (["content writer", "copywriter", "technical writer"],
         ["Build a writing portfolio/blog", "Learn SEO writing basics", "Show range across formats (blogs, ads, docs)"]),
        (["teacher", "professor", "lecturer", "educator"],
         ["Highlight curriculum design experience", "Add student outcome/result improvements", "Get certified in relevant teaching tools"]),
        (["civil engineer"],
         ["Get certified in AutoCAD/Revit", "Add site project experience", "Learn structural analysis software"]),
        (["mechanical engineer"],
         ["Get certified in SolidWorks/AutoCAD", "Add hands-on design/manufacturing projects", "Learn GD&T basics"]),
        (["electrical engineer"],
         ["Get certified in circuit design tools", "Add PCB or embedded systems projects", "Learn relevant safety standards"]),
        (["nurse", "healthcare", "medical", "doctor"],
         ["Highlight patient care outcomes", "Add relevant certifications (BLS/ACLS etc.)", "Show experience with medical software/EHR"]),
        (["customer support", "customer service", "call center"],
         ["Highlight resolution time/satisfaction metrics", "Add CRM/helpdesk tool experience", "Show examples of handling difficult cases"]),
    ]
    for keywords, recs in role_data:
        if any(k in role_lower for k in keywords):
            return recs
    return ["Keep improving your resume", "Add measurable achievements", "Tailor your resume to the job description"]

def generate_interview_questions(job_role):
    role_lower = job_role.lower()
    role_data = [
        (["data analyst", "data analytics", "business intelligence", "bi analyst"],
         ["Explain normalization in SQL", "How do you handle missing data?"]),
        (["software engineer", "software developer", "backend", "full stack", "sde"],
         ["What is polymorphism?", "Explain REST APIs"]),
        (["ai engineer", "machine learning", "ml engineer", "data scientist"],
         ["Difference between supervised and unsupervised learning?", "Explain transformers in NLP"]),
        (["web developer", "frontend", "front-end", "react", "javascript developer"],
         ["What is the virtual DOM?", "Explain the difference between let, const, and var"]),
        (["mobile app developer", "android developer", "ios developer", "flutter"],
         ["Explain the activity/view lifecycle", "How do you manage app state?"]),
        (["marketing", "seo", "content", "digital marketing"],
         ["How do you measure campaign ROI?", "Explain on-page vs off-page SEO"]),
        (["hr", "human resource", "recruit", "talent"],
         ["How do you handle a conflict between employees?", "What is your approach to talent sourcing?"]),
        (["devops", "cloud engineer", "site reliability", "sre"],
         ["Explain CI/CD pipeline", "What is the difference between containers and VMs?"]),
        (["cybersecurity", "security analyst", "ethical hacker", "penetration tester"],
         ["Explain the difference between symmetric and asymmetric encryption", "What is a SQL injection attack?"]),
        (["qa", "quality assurance", "software tester", "test engineer"],
         ["Explain the difference between smoke and regression testing", "How do you write an effective test case?"]),
        (["business analyst", "product analyst"],
         ["How do you gather requirements from stakeholders?", "Explain the difference between BRD and FRD"]),
        (["project manager", "program manager", "scrum master"],
         ["How do you handle scope creep?", "Explain the difference between Agile and Waterfall"]),
        (["accountant", "finance", "financial analyst"],
         ["Explain the difference between accounts payable and receivable", "How do you ensure accuracy in financial reports?"]),
        (["sales", "business development"],
         ["How do you handle a difficult client?", "Walk me through your sales process"]),
        (["graphic designer", "ui/ux designer", "ux designer", "product designer"],
         ["Walk me through your design process", "How do you handle design feedback you disagree with?"]),
        (["content writer", "copywriter", "technical writer"],
         ["How do you adapt tone for different audiences?", "How do you research a topic you're unfamiliar with?"]),
        (["teacher", "professor", "lecturer", "educator"],
         ["How do you handle students with different learning paces?", "How do you assess student understanding?"]),
        (["civil engineer"],
         ["Explain the difference between dead load and live load", "How do you ensure quality control on site?"]),
        (["mechanical engineer"],
         ["Explain the difference between stress and strain", "How do you approach a design failure analysis?"]),
        (["electrical engineer"],
         ["Explain the difference between AC and DC", "How do you approach circuit troubleshooting?"]),
        (["nurse", "healthcare", "medical", "doctor"],
         ["How do you handle a high-pressure emergency situation?", "How do you ensure patient safety and accuracy?"]),
        (["customer support", "customer service", "call center"],
         ["How do you handle an angry customer?", "How do you prioritize multiple customer issues at once?"]),
    ]
    for keywords, qs in role_data:
        if any(k in role_lower for k in keywords):
            return qs
    return ["Tell me about yourself", "Why should we hire you for this role?"]

def predict_salary(job_role):
    role_lower = job_role.lower()
    role_data = [
        (["data analyst", "data analytics", "business intelligence", "bi analyst"], "Rs. 5-8 LPA"),
        (["software engineer", "software developer", "backend", "full stack", "sde"], "Rs. 6-12 LPA"),
        (["ai engineer", "machine learning", "ml engineer", "data scientist"], "Rs. 10-20 LPA"),
        (["web developer", "frontend", "front-end", "react", "javascript developer"], "Rs. 4-9 LPA"),
        (["mobile app developer", "android developer", "ios developer", "flutter"], "Rs. 5-10 LPA"),
        (["marketing", "seo", "content", "digital marketing"], "Rs. 4-8 LPA"),
        (["hr", "human resource", "recruit", "talent"], "Rs. 4-7 LPA"),
        (["devops", "cloud engineer", "site reliability", "sre"], "Rs. 8-15 LPA"),
        (["cybersecurity", "security analyst", "ethical hacker", "penetration tester"], "Rs. 6-14 LPA"),
        (["qa", "quality assurance", "software tester", "test engineer"], "Rs. 4-8 LPA"),
        (["business analyst", "product analyst"], "Rs. 6-10 LPA"),
        (["project manager", "program manager", "scrum master"], "Rs. 10-18 LPA"),
        (["accountant", "finance", "financial analyst"], "Rs. 4-9 LPA"),
        (["sales", "business development"], "Rs. 3-8 LPA"),
        (["graphic designer", "ui/ux designer", "ux designer", "product designer"], "Rs. 4-9 LPA"),
        (["content writer", "copywriter", "technical writer"], "Rs. 3-6 LPA"),
        (["teacher", "professor", "lecturer", "educator"], "Rs. 3-7 LPA"),
        (["civil engineer"], "Rs. 4-8 LPA"),
        (["mechanical engineer"], "Rs. 4-8 LPA"),
        (["electrical engineer"], "Rs. 4-8 LPA"),
        (["nurse", "healthcare", "medical", "doctor"], "Rs. 4-10 LPA"),
        (["customer support", "customer service", "call center"], "Rs. 3-6 LPA"),
    ]
    for keywords, salary in role_data:
        if any(k in role_lower for k in keywords):
            return salary
    return "Rs. 5-10 LPA"

def export_pdf_report(name, job_role, ats_score, match_score, recommendations, interview_questions, salary_prediction):
    pdf = FPDF()
    pdf.add_page()

    if UNICODE_FONT_AVAILABLE:
        # Unicode font found - supports any special characters (e.g. rupee symbol, accents)
        pdf.add_font("DejaVu", "", FONT_PATH, uni=True)
        pdf.set_font("DejaVu", size=12)
    else:
        # Fall back to a built-in core font so the app doesn't crash.
        # Core fonts (Arial/Helvetica) only support latin-1, so make sure
        # predict_salary() etc. don't contain characters like ₹.
        pdf.set_font("Arial", size=12)

    effective_width = pdf.w - pdf.l_margin - pdf.r_margin

    pdf.cell(200, 10, f"Resume Analysis Report for {name}", ln=True, align="C")

    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(effective_width, 10, f"Job Role: {job_role}\nATS Score: {ats_score}%\nMatch Score: {match_score}%\n\nRecommendations:\n" + "\n".join(recommendations))

    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(effective_width, 10, "\nInterview Questions:\n" + "\n".join(interview_questions))

    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(effective_width, 10, f"\nSalary Prediction: {salary_prediction}")
    pdf.output("resume_report.pdf")

def export_word_report(name, job_role, ats_score, match_score, recommendations, interview_questions, salary_prediction):
    doc = Document()
    doc.add_heading(f"Resume Analysis Report for {name}", 0)
    doc.add_paragraph(f"Job Role: {job_role}")
    doc.add_paragraph(f"ATS Score: {ats_score}%")
    doc.add_paragraph(f"Match Score: {match_score}%")
    doc.add_heading("Recommendations", level=1)
    for rec in recommendations:
        doc.add_paragraph(rec, style="List Bullet")
    doc.add_heading("Interview Questions", level=1)
    for q in interview_questions:
        doc.add_paragraph(q, style="List Bullet")
    doc.add_paragraph(f"Salary Prediction: {salary_prediction}")
    doc.save("resume_report.docx")

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="AI Resume Analyzer", layout="wide")
st.title("📄 AI Resume Analyzer Dashboard")

name = st.text_input("Candidate Name")

job_role_options = ["Data Analyst", "Software Engineer", "AI Engineer", "Other (type manually)"]
job_role_choice = st.selectbox("Select Job Role", job_role_options)

if job_role_choice == "Other (type manually)":
    job_role = st.text_input("Type your job role")
else:
    job_role = job_role_choice

uploaded_resume = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
job_description = st.text_area("Paste Job Description")

# Default JD fallback
default_jds = {
    "Data Analyst": "Looking for a Data Analyst with skills in SQL, Python, Tableau, and statistics.",
    "Software Engineer": "Seeking a Software Engineer with experience in Java, C++, system design, and algorithms.",
    "AI Engineer": "Hiring AI Engineer with expertise in machine learning, NLP, and TensorFlow/PyTorch."
}
if not job_description.strip():
    job_description = default_jds.get(job_role, f"Looking for a {job_role} with relevant skills and experience.")

if st.button("Analyze Resume"):
    if uploaded_resume and name and job_role.strip():
        # Extract resume text
        if uploaded_resume.type == "application/pdf":
            resume_text = extract_text_from_pdf(uploaded_resume)
        else:
            resume_text = extract_text_from_docx(uploaded_resume)

        # ATS Score
        ats_score = calculate_ats_score(resume_text, job_description)
        match_score = ats_score  # simplified

        # Recommendations, Interview Qs, Salary
        recommendations = generate_recommendations(job_role)
        interview_questions = generate_interview_questions(job_role)
        salary_prediction = predict_salary(job_role)

        # Save to DB
        c.execute("INSERT INTO resumes (name, job_role, ats_score, match_score, recommendations, interview_questions, salary_prediction) VALUES (?, ?, ?, ?, ?, ?, ?)",
                  (name, job_role, ats_score, match_score, ", ".join(recommendations), ", ".join(interview_questions), salary_prediction))
        conn.commit()

        # Dashboard
        st.subheader("📊 Analysis Results")
        st.write(f"**ATS Score:** {ats_score}%")
        st.write(f"**Match Score:** {match_score}%")
        st.write(f"**Salary Prediction:** {salary_prediction}")

        # Charts
        fig, ax = plt.subplots()
        ax.bar(["ATS Score", "Match Score"], [ats_score, match_score], color=["blue", "green"])
        st.pyplot(fig)

        # Recommendations
        st.subheader("✅ Recommendations")
        for rec in recommendations:
            st.write("- " + rec)

        # Interview Questions
        st.subheader("🎤 Interview Questions")
        for q in interview_questions:
            st.write("- " + q)

        # Export Reports with Download Buttons
        export_pdf_report(name, job_role, ats_score, match_score, recommendations, interview_questions, salary_prediction)
        with open("resume_report.pdf", "rb") as f:
            st.download_button("📥 Download PDF Report", f, file_name="resume_report.pdf")

        export_word_report(name, job_role, ats_score, match_score, recommendations, interview_questions, salary_prediction)
        with open("resume_report.docx", "rb") as f:
            st.download_button("📥 Download Word Report", f, file_name="resume_report.docx")
